import os
from docx import Document
import random
import re
from datetime import datetime

def extract_year(text):
    """从文本中提取年份"""
    year_match = re.search(r'(\d{4})', text)
    return int(year_match.group(1)) if year_match else 9999

def classify_subject(text):
    """对书籍进行学科分类"""
    subjects = {
        '力学': ['力学', '静力学', '动力学'],
        '材料': ['材料', '耐力'],
        '建筑': ['建筑', '构造', '钢筋', '混凝土'],
    }
    
    text_lower = text.lower()
    for category, keywords in subjects.items():
        if any(keyword in text_lower for keyword in keywords):
            return category
    return '其他'

def format_book_info(text):
    """格式化书籍信息"""
    # 处理作者信息
    author_match = re.search(r'\[(.*?)\]', text)
    if author_match:
        author = author_match.group(1)
        text = text.replace(f'[{author}]', '')
    else:
        # 查找其他形式的作者标记
        author_match = re.search(r'[，,](.*?)编[著]?', text)
        author = author_match.group(1) if author_match else ''
    
    # 清理和标准化书名
    title = text.strip()
    title = re.sub(r'\s+', ' ', title)
    
    return {'title': title, 'author': author.strip()}

def read_doc_files(directory):
    """读取目录下所有doc文件并进行分类整理"""
    if not os.path.exists(directory):
        print(f"错误：目录 '{directory}' 不存在")
        return {}
        
    books_data = {}
    doc_files = [f for f in os.listdir(directory) if f.endswith(('.doc', '.docx'))]
    
    if not doc_files:
        print(f"警告：在 '{directory}' 目录中没有找到.doc或.docx文件")
        return {}
        
    for filename in doc_files:
        try:
            filepath = os.path.join(directory, filename)
            doc = Document(filepath)
            if len(doc.paragraphs) < 2:
                print(f"警告：文件 '{filename}' 格式不正确，需要至少包含两个段落")
                continue
                
            book_name = doc.paragraphs[0].text.strip()
            classification = doc.paragraphs[1].text.strip().split('/')
            
            # 提取和处理书籍信息
            book_info = format_book_info(book_name)
            year = extract_year(book_name)
            subject = classify_subject(book_name)
            
            # 组织数据结构
            if subject not in books_data:
                books_data[subject] = {}
            
            decade = (year // 10) * 10
            if decade not in books_data[subject]:
                books_data[subject][decade] = []
                
            books_data[subject][decade].append({
                'title': book_info['title'],
                'author': book_info['author'],
                'year': year,
                'classification': classification
            })
            
            print(f"成功读取文件：{filename}")
        except Exception as e:
            print(f"处理文件 '{filename}' 时出错：{str(e)}")
            
    return books_data

def generate_distance():
    """生成随机的进化距离"""
    return round(random.uniform(0.1, 0.5), 2)  # 缩小距离范围，使树更紧凑

def build_tree(books_data, indent=True, show_distance=True):
    """构建优化后的系统发育树"""
    def convert_to_newick(node, name=None, level=0):
        if isinstance(node, list):  # 叶子节点（书籍列表）
            children = []
            for book in sorted(node, key=lambda x: x['year']):
                book_name = f"{book['title']}"
                if book['author']:
                    book_name += f" ({book['author']})"
                if book['year'] != 9999:
                    book_name += f" [{book['year']}]"
                    
                distance = f":{generate_distance()}" if show_distance else ""
                children.append(f"{book_name}{distance}")
                
            if indent:
                indent_str = "\n" + "  " * level
                return f"({indent_str}{(','+indent_str).join(children)}{indent_str})"
            return f"({','.join(children)})"
            
        # 非叶子节点（分类层级）
        children = []
        for child_name, child_node in sorted(node.items()):
            child_str = convert_to_newick(child_node, child_name, level + 1)
            if child_str:  # 只添加非空的子树
                children.append(child_str)
                
        if not children:
            return ""
            
        if indent:
            indent_str = "\n" + "  " * level
            result = f"({indent_str}{(','+indent_str).join(children)}{indent_str})"
        else:
            result = f"({','.join(children)})"
            
        if name:
            distance = f":{generate_distance()}" if show_distance else ""
            result += f"{name}{distance}"
        return result

    tree = convert_to_newick(books_data) + ";"
    return tree

def main():
    directory = "data"
    books_data = read_doc_files(directory)
    
    if not books_data:
        print("没有可用的书籍数据，无法生成系统发育树")
        return
    
    # 生成两种格式的树
    tree_indented = build_tree(books_data, indent=True, show_distance=True)
    tree_compact = build_tree(books_data, indent=False, show_distance=True)
    
    try:
        # 保存缩进格式的树
        with open("phylogenetic_tree_indented.txt", "w", encoding="utf-8") as f:
            f.write(tree_indented)
        print("缩进格式的系统发育树已保存到 phylogenetic_tree_indented.txt")
        
        # 同时生成一个更易读的文本报告
        with open("books_report.txt", "w", encoding="utf-8") as f:
            f.write("力学书籍分类报告\n")
            f.write("=" * 50 + "\n\n")
            
            for subject, decades in sorted(books_data.items()):
                f.write(f"\n{subject}类书籍：\n")
                f.write("-" * 30 + "\n")
                
                for decade, books in sorted(decades.items()):
                    if decade == 9990:
                        f.write("\n年代未知：\n")
                    else:
                        f.write(f"\n{decade}年代：\n")
                        
                    for book in sorted(books, key=lambda x: x['year']):
                        book_info = f"- {book['title']}"
                        if book['author']:
                            book_info += f" (作者：{book['author']})"
                        if book['year'] != 9999:
                            book_info += f" [{book['year']}]"
                        f.write(book_info + "\n")
                
        print("已生成易读的分类报告：books_report.txt")
        
    except Exception as e:
        print(f"保存文件时出错：{str(e)}")

if __name__ == "__main__":
    main()
